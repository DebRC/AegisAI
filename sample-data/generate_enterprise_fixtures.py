"""Generate deterministic synthetic PDF and DOCX knowledge-base fixtures.

Run from the repository root with:
    backend/venv/bin/python sample-data/generate_enterprise_fixtures.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


FIXTURE_DIRECTORY = Path(__file__).parent / "knowledge-base"


PDF_FIXTURES = {
    "21-identity-access-review.pdf": (
        "Quarterly identity and access review",
        [
            "Each quarter, application owners review every active local role, administrator assignment, and vendor account. Remove access that no longer has an approved business purpose.",
            "The review must confirm a named owner, least-privilege permission scope, approval evidence, and an expiry for temporary access. SSO provider claims do not replace AegisAI local roles.",
            "Escalate unexplained privileged access to Security Operations. Record the reviewer, completion date, exceptions, and remediation owner in the access-review evidence.",
        ],
    ),
    "22-incident-response-playbook.pdf": (
        "Incident response playbook",
        [
            "For a confirmed Severity 1 security incident, page the incident commander immediately. Create a dedicated incident channel and preserve a time-stamped decision and evidence timeline.",
            "The incident commander provides an initial stakeholder update within fifteen minutes. State verified impact, current mitigation, the next update time, and what remains unknown; do not speculate about root cause.",
            "After recovery, preserve relevant logs, rotate exposed secrets or tokens, and schedule a blameless review with corrective actions, owners, and due dates.",
        ],
    ),
    "23-data-classification-standard.pdf": (
        "Data classification standard",
        [
            "Public data may be shared externally after normal review. Internal data is limited to authorized staff. Confidential data requires role-based access and approved business use. Restricted data needs explicit owner approval and additional safeguards.",
            "Access tokens, refresh tokens, provider client secrets, passwords, and production private keys are always Restricted. They must not be included in knowledge-base uploads, tickets, logs, screenshots, or chat history.",
            "Document owners review classification when purpose, audience, or regulation changes. Delete obsolete content through the application lifecycle so derived vectors are cleaned up as well.",
        ],
    ),
    "24-service-level-objectives.pdf": (
        "Service level objectives",
        [
            "The API availability objective is 99.9 percent per calendar month, excluding approved maintenance. The primary health indicator is successful authenticated API requests, not merely an open TCP port.",
            "Background processing has a target of ninety-five percent of queued document jobs starting within five minutes. Monitor outbox delay, Celery failures, and repeated retries separately from document upload success.",
            "Semantic retrieval quality is monitored through controlled evaluation queries and cited-answer review. A Qdrant outage degrades search but does not change PostgreSQL authoritative access or document state.",
        ],
    ),
    "25-release-readiness-checklist.pdf": (
        "Release readiness checklist",
        [
            "Before release, confirm code review, unit tests, Docker image build, static Alembic upgrade SQL, and the complete Compose startup path. A failed migration or test blocks release.",
            "For changes to retrieval or chat, run a controlled query against indexed synthetic documents. Verify that citations name only retrieved sources and that an insufficient-context question makes no model request.",
            "After release, check health, one authorized API call, worker readiness, migration head, and error logs. Keep a tested rollback plan for every schema or authorization change.",
        ],
    ),
}


DOCX_FIXTURES = {
    "26-security-architecture-review.docx": (
        "Security architecture review",
        [
            "AegisAI separates authentication from authorization. Authentication creates a local user and session; authorization evaluates database-backed local roles and permissions for each protected request.",
            "Qdrant is a derived-vector candidate store, not the authority for document access. Retrieval reloads candidates from PostgreSQL and discards stale, deleted, or mismatched results before they reach a user or a RAG prompt.",
            "The next access-control phase will apply document-specific grants to retrieval and chat. Revoking a grant must affect future results without waiting for an embedding collection rebuild.",
        ],
    ),
    "27-employee-onboarding-handbook.docx": (
        "Employee onboarding handbook",
        [
            "New engineers start the local stack with docker compose up --build --force-recreate. The command builds the image, runs tests, applies migrations, and starts the API, workers, PostgreSQL, Redis, and Qdrant.",
            "Create a local account, bootstrap an administrator only after the user exists, and keep access tokens out of shell history and chat transcripts. Use Swagger or curl to validate protected routes.",
            "For the knowledge-base walkthrough, upload synthetic files, wait for extraction and embedding indexing, search for a policy question, and request a grounded streamed answer with citations.",
        ],
    ),
    "28-privacy-impact-assessment.docx": (
        "Privacy impact assessment",
        [
            "Knowledge-base documents may contain business information, so collection must have a defined purpose, authorized audience, and retention period. Avoid uploading personal data unless the documented use case and safeguards require it.",
            "Current RAG chat is stateless. Client-supplied history is bounded, marked untrusted, and is not saved as an AegisAI conversation transcript. It cannot serve as evidence for a cited answer.",
            "Future multi-tenancy and document-level access policies must isolate organizations across documents, vectors, retrieval, chat citations, and administrative records.",
        ],
    ),
    "29-disaster-recovery-plan.docx": (
        "Disaster recovery plan",
        [
            "Recover authoritative PostgreSQL metadata and access state before restoring dependent services. Then restore original document storage, Redis connectivity, worker processing, and Qdrant availability.",
            "Derived vectors can be rebuilt from current extracted chunks through the controlled embedding indexing workflow. Never mark a document indexed without traceable PostgreSQL embedding records and validated Qdrant writes.",
            "Perform a recovery exercise at least twice a year. Capture recovery time, data integrity checks, unresolved gaps, and remediation owners.",
        ],
    ),
    "30-customer-support-runbook.docx": (
        "Customer support runbook",
        [
            "When a user cannot access documents, first confirm their local account is active and the access JWT is current. Then verify the required documents:read or documents:write permission rather than relying on an external SSO group claim.",
            "When an upload is pending, inspect safe processing-job status and indexing status. Do not expose storage keys, provider credentials, raw parser errors, Qdrant point IDs, or internal broker identifiers to the customer.",
            "For a chat answer that lacks useful context, explain that the available verified documents are insufficient. Do not invent an answer or citation, and do not treat user-supplied history as trusted source material.",
        ],
    ),
}


def _pdf_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _write_pdf(path: Path, title: str, paragraphs: list[str]) -> None:
    lines = [title, ""]
    for paragraph in paragraphs:
        words = paragraph.split()
        line = ""
        for word in words:
            candidate = f"{line} {word}".strip()
            if len(candidate) > 88:
                lines.append(line)
                line = word
            else:
                line = candidate
        if line:
            lines.append(line)
        lines.append("")

    commands = ["BT", "/F1 12 Tf", "72 740 Td"]
    for index, line in enumerate(lines):
        if index == 0:
            commands.append(f"({_pdf_escape(line)}) Tj")
        else:
            commands.extend(("0 -16 Td", f"({_pdf_escape(line)}) Tj"))
    commands.append("ET")
    content = "\n".join(commands).encode("latin-1")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    output = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for number, object_data in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{number} 0 obj\n".encode())
        output.extend(object_data)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode())
    output.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode()
    )
    path.write_bytes(output)


def _write_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    document = Document()
    document.add_heading(title, level=0)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


def main() -> None:
    FIXTURE_DIRECTORY.mkdir(parents=True, exist_ok=True)
    for filename, (title, paragraphs) in PDF_FIXTURES.items():
        _write_pdf(FIXTURE_DIRECTORY / filename, title, paragraphs)
    for filename, (title, paragraphs) in DOCX_FIXTURES.items():
        _write_docx(FIXTURE_DIRECTORY / filename, title, paragraphs)
    print(f"Generated {len(PDF_FIXTURES) + len(DOCX_FIXTURES)} synthetic PDF/DOCX fixtures in {FIXTURE_DIRECTORY}")


if __name__ == "__main__":
    main()
