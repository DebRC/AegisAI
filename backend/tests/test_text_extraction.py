from io import BytesIO
import importlib.util
import unittest

from pydantic import ValidationError

from app.core.config import Settings
from app.extraction.base import ExtractedText, ExtractedTextBlock, SourceLocation
from app.extraction.exceptions import EncryptedDocumentError
from app.extraction.exceptions import ExtractedTextLimitExceededError
from app.extraction.exceptions import NoExtractableTextError
from app.extraction.exceptions import TextDecodingError
from app.extraction.exceptions import TextExtractionProviderError
from app.extraction.exceptions import UnsupportedDocumentTypeError
from app.extraction.extractors import DocxTextExtractor, PdfTextExtractor, Utf8TextExtractor
from app.extraction.registry import TextExtractorRegistry


class StaticExtractor:
    def __init__(self, text: str):
        self.text = text

    def extract(self, source: BytesIO) -> ExtractedText:
        return ExtractedText(blocks=(ExtractedTextBlock(self.text),))


class TextExtractionValueTests(unittest.TestCase):
    def test_source_location_and_text_blocks_validate_input_and_preserve_order(self) -> None:
        first = ExtractedTextBlock("First", SourceLocation("page", 1))
        second = ExtractedTextBlock("Second", SourceLocation("page", 2))

        result = ExtractedText(blocks=(first, second))

        self.assertEqual(result.text, "First\n\nSecond")
        with self.assertRaises(ValueError):
            SourceLocation("", 1)
        with self.assertRaises(ValueError):
            SourceLocation("page", 0)
        with self.assertRaises(ValueError):
            ExtractedTextBlock("   ")
        with self.assertRaises(ValueError):
            ExtractedText(blocks=())


class Utf8TextExtractorTests(unittest.TestCase):
    def setUp(self) -> None:
        self.extractor = Utf8TextExtractor()

    def test_decodes_utf8_with_bom_and_keeps_text_as_one_document_block(self) -> None:
        result = self.extractor.extract(BytesIO("\ufeff# AegisAI\nPolicy".encode()))

        self.assertEqual(result.text, "# AegisAI\nPolicy")
        self.assertEqual(result.blocks[0].source_location, SourceLocation("document"))

    def test_rejects_invalid_utf8_and_empty_text(self) -> None:
        with self.assertRaises(TextDecodingError):
            self.extractor.extract(BytesIO(b"\xff\xfe"))
        with self.assertRaises(NoExtractableTextError):
            self.extractor.extract(BytesIO(b" \n\t"))


class PdfTextExtractorTests(unittest.TestCase):
    def test_extracts_nonempty_pages_with_one_based_page_locations(self) -> None:
        class Page:
            def __init__(self, text: str | None):
                self.text = text

            def extract_text(self) -> str | None:
                return self.text

        class Reader:
            is_encrypted = False
            pages = [Page("First page"), Page(""), Page("Third page")]

        result = PdfTextExtractor(reader_factory=lambda source: Reader()).extract(BytesIO(b"pdf"))

        self.assertEqual(result.text, "First page\n\nThird page")
        self.assertEqual(result.blocks[1].source_location, SourceLocation("page", 3))

    def test_rejects_encrypted_empty_and_unreadable_pdf_sources(self) -> None:
        class EncryptedReader:
            is_encrypted = True
            pages: list[object] = []

        class EmptyReader:
            is_encrypted = False
            pages: list[object] = []

        with self.assertRaises(EncryptedDocumentError):
            PdfTextExtractor(reader_factory=lambda source: EncryptedReader()).extract(BytesIO(b"pdf"))
        with self.assertRaises(NoExtractableTextError):
            PdfTextExtractor(reader_factory=lambda source: EmptyReader()).extract(BytesIO(b"pdf"))
        with self.assertRaises(TextExtractionProviderError):
            PdfTextExtractor(reader_factory=lambda source: (_ for _ in ()).throw(ValueError())).extract(BytesIO(b"pdf"))


class DocxTextExtractorTests(unittest.TestCase):
    def test_extracts_nonempty_paragraphs_with_one_based_locations(self) -> None:
        class Paragraph:
            def __init__(self, text: str):
                self.text = text

        class DocxFile:
            paragraphs = [Paragraph("Introduction"), Paragraph(""), Paragraph("Controls")]

        result = DocxTextExtractor(document_factory=lambda source: DocxFile()).extract(BytesIO(b"docx"))

        self.assertEqual(result.text, "Introduction\n\nControls")
        self.assertEqual(result.blocks[1].source_location, SourceLocation("paragraph", 3))

    def test_maps_empty_or_unreadable_docx_to_safe_errors(self) -> None:
        class EmptyDocxFile:
            paragraphs: list[object] = []

        with self.assertRaises(NoExtractableTextError):
            DocxTextExtractor(document_factory=lambda source: EmptyDocxFile()).extract(BytesIO(b"docx"))
        with self.assertRaises(TextExtractionProviderError):
            DocxTextExtractor(document_factory=lambda source: (_ for _ in ()).throw(ValueError())).extract(BytesIO(b"docx"))


class TextExtractorRegistryTests(unittest.TestCase):
    def test_dispatches_normalized_allowed_mime_types_and_enforces_limit(self) -> None:
        registry = TextExtractorRegistry(
            10,
            plain_text_extractor=StaticExtractor("policy"),
            pdf_extractor=StaticExtractor("pdf text"),
            docx_extractor=StaticExtractor("docx text"),
        )

        self.assertEqual(registry.extract(content_type="text/plain; charset=utf-8", source=BytesIO()).text, "policy")
        self.assertEqual(registry.extract(content_type="application/pdf", source=BytesIO()).text, "pdf text")
        self.assertEqual(
            registry.extract(
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                source=BytesIO(),
            ).text,
            "docx text",
        )
        with self.assertRaises(UnsupportedDocumentTypeError):
            registry.extract(content_type="application/octet-stream", source=BytesIO())
        with self.assertRaises(UnsupportedDocumentTypeError):
            registry.extract(content_type=object(), source=BytesIO())
        with self.assertRaises(ExtractedTextLimitExceededError):
            TextExtractorRegistry(5, plain_text_extractor=StaticExtractor("sixsix")).extract(
                content_type="text/markdown",
                source=BytesIO(),
            )

    def test_rejects_invalid_character_limits(self) -> None:
        for value in (0, -1, True, "100"):
            with self.subTest(value=value):
                with self.assertRaises(ValueError):
                    TextExtractorRegistry(value)


class TextExtractionConfigurationTests(unittest.TestCase):
    _REQUIRED_SETTINGS = {
        "APP_NAME": "AegisAI",
        "APP_VERSION": "test",
        "APP_ENV": "test",
        "HOST": "127.0.0.1",
        "PORT": 8000,
        "DATABASE_URL": "sqlite+pysqlite:///:memory:",
        "QDRANT_URL": "http://qdrant:6333",
        "JWT_SECRET_KEY": "test-secret",
        "JWT_ALGORITHM": "HS256",
        "ACCESS_TOKEN_EXPIRE_MINUTES": 15,
        "REFRESH_TOKEN_EXPIRE_DAYS": 7,
    }

    def test_uses_safe_default_extraction_and_chunking_limits(self) -> None:
        settings = Settings(**self._REQUIRED_SETTINGS)

        self.assertEqual(settings.DOCUMENT_MAX_EXTRACTED_TEXT_CHARACTERS, 5_000_000)
        self.assertEqual(settings.DOCUMENT_CHUNK_TARGET_CHARACTERS, 1_200)
        self.assertEqual(settings.DOCUMENT_CHUNK_OVERLAP_CHARACTERS, 200)

    def test_rejects_non_positive_or_invalid_chunking_configuration(self) -> None:
        with self.assertRaises(ValidationError):
            Settings(**self._REQUIRED_SETTINGS, DOCUMENT_MAX_EXTRACTED_TEXT_CHARACTERS=0)
        with self.assertRaises(ValidationError):
            Settings(
                **self._REQUIRED_SETTINGS,
                DOCUMENT_CHUNK_TARGET_CHARACTERS=200,
                DOCUMENT_CHUNK_OVERLAP_CHARACTERS=200,
            )


@unittest.skipUnless(importlib.util.find_spec("docx"), "python-docx is installed in the Docker image")
class RealDocxExtractorTests(unittest.TestCase):
    def test_loads_a_real_docx_stream(self) -> None:
        from docx import Document as DocxDocument

        source = BytesIO()
        document = DocxDocument()
        document.add_paragraph("Introduction")
        document.add_paragraph("Controls")
        document.save(source)
        source.seek(0)

        result = DocxTextExtractor().extract(source)

        self.assertEqual(result.text, "Introduction\n\nControls")


@unittest.skipUnless(importlib.util.find_spec("pypdf"), "pypdf is installed in the Docker image")
class RealPdfExtractorTests(unittest.TestCase):
    def test_loads_a_real_pdf_stream(self) -> None:
        source = BytesIO(_pdf_with_text("AegisAI policy"))

        result = PdfTextExtractor().extract(source)

        self.assertIn("AegisAI policy", result.text)
        self.assertEqual(result.blocks[0].source_location, SourceLocation("page", 1))


def _pdf_with_text(text: str) -> bytes:
    """Build the smallest valid single-page PDF needed for adapter testing."""
    content = f"BT /F1 18 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(content)).encode("ascii") + b" >>\nstream\n" + content + b"\nendstream",
    ]
    document = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for object_number, object_value in enumerate(objects, start=1):
        offsets.append(len(document))
        document.extend(f"{object_number} 0 obj\n".encode("ascii"))
        document.extend(object_value)
        document.extend(b"\nendobj\n")
    xref_offset = len(document)
    document.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    document.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        document.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    document.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode(
            "ascii"
        )
    )
    return bytes(document)
